import tkinter as tk
from tkinter import messagebox, filedialog
from openpyxl import Workbook, load_workbook  # Para Excel
from docx import Document  # Para Word
import pandas as pd
from datetime import datetime, timedelta
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from tkcalendar import Calendar
import matplotlib.pyplot as plt
from docx.oxml.ns import qn
import unicodedata

def normalizar_texto(texto):
    texto = str(texto).strip()
    texto = unicodedata.normalize('NFD', texto).encode('ascii', 'ignore').decode('utf-8')
    texto = texto.lower()  # Convertir a minÃºsculas
    return texto

# FunciÃ³n para cambiar al menÃº principal
def cambiar_a_menu():
    if frame in frame_anexo:
        for frame in [frame_anexo]:
            frame.pack_forget()    
    frame_menu.pack(fill="both", expand=True)

def frame_semanal():
    frame_menu.pack_forget()
    frame_semanal.pack(fill="both", expand=True)

def cambiar_frame(actual, siguiente):
    actual.pack_forget()
    siguiente.pack(fill="both", expand=True)

# Funciones para cambiar el estilo del botÃ³n cuando el mouse pasa sobre Ã©l
def on_enter(event):
    event.widget.config(bg="#4c4c4c", fg="white")

def on_leave(event):
    event.widget.config(bg="#dbdcdc", fg="black")    

# FunciÃ³n para cargar el archivo Excel
def archivo_anexo():
    global data # Declarar la variable global
    # Seleccionar archivo Excel
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if not file_path:
        return

    try:
        # Leer el archivo
        data = pd.read_excel(file_path, skiprows=4)

        # Limpiar nombres de columnas
        data.columns = data.columns.str.strip().str.replace(r'\s+', ' ', regex=True)

        # Validar columnas obligatorias
        columnas_requeridas = [
            "Estado MMS", "SISTEMA", "TIPO", "NOMBRE DEL EQUIPO",
            "OTROS SISTEMAS", "EMPLAZAMIENTO", "LINEA", "TRATAMIENTO"
        ]
        columnas_faltantes = [col for col in columnas_requeridas if col not in data.columns]
        if columnas_faltantes:
            messagebox.showerror("Error", f"Columnas faltantes en el archivo: {', '.join(columnas_faltantes)}")
            return
        
        # Convertir la columna "Fecha" a formato datetime
        if "Fecha" in data.columns:
            data["Fecha"] = pd.to_datetime(data["Fecha"], errors="coerce")  # Convertir a datetime
            data = data.dropna(subset=["Fecha"])  # Eliminar filas con fechas invÃ¡lidas

            # Extraer los aÃ±os Ãºnicos
            anios_disponibles = data["Fecha"].dt.year.dropna().astype(int).unique().tolist()
            anios_disponibles.sort()
            anios_disponibles.insert(0, "TODOS")  # Agregar opciÃ³n "TODOS"


            # Actualizar el menÃº desplegable para los aÃ±os
            if "Fecha" not in menus:
                label_fecha = tk.Label(frame_campos, text="Fecha (AÃ±o)", bg="#2b2b2b", fg="white", font=("arial", 10))
                label_fecha.grid(row=0, column=0, padx=10, pady=5, sticky="e")
                menus["Fecha"] = tk.OptionMenu(frame_campos, variables["Fecha"], *anios_disponibles, command=lambda v: actualizar_seleccion("Fecha", v))
                menus["Fecha"].grid(row=0, column=1, padx=10, pady=5, sticky="w")
            else:
                menus["Fecha"].children["menu"].delete(0, "end")
                for anio in anios_disponibles:
                    menus["Fecha"].children["menu"].add_command(label=anio, command=lambda v=anio: actualizar_seleccion("Fecha", v))

    except Exception as e:
        messagebox.showerror("Error", f"Error al cargar el archivo: {e}")


        # Rellenar valores nulos con un texto predeterminado
        for columna in columnas_requeridas:
            data[columna] = data[columna].fillna("Desconocido").astype(str).str.strip()

    except Exception as e:
        messagebox.showerror("Error", f"Error al cargar el archivo: {e}")

    # Actualizar semanas disponibles (si aplicable)
    if "Fecha" in data.columns:
        data["Semana"] = data["Fecha"].dt.isocalendar().week
        semanas_disponibles = data["Semana"].dropna().astype(int).unique().tolist()
        semanas_disponibles.sort()
        semanas_disponibles.insert(0, "TODOS")  # Agregar opciÃ³n "TODOS"

        if "Semana" not in menus:
            label_semana = tk.Label(frame_campos, text="Semana", bg="#2b2b2b", fg="white", font=("arial", 10))
            label_semana.grid(row=1, column=0, padx=10, pady=5, sticky="e")
            menus["Semana"] = tk.OptionMenu(frame_campos, variables["Semana"], *semanas_disponibles, command=lambda v: actualizar_seleccion("Semana", v))
            menus["Semana"].grid(row=1, column=1, padx=10, pady=5, sticky="w")
        else:
            menus["Semana"].children["menu"].delete(0, "end")
            for semana in semanas_disponibles:
                menus["Semana"].children["menu"].add_command(
                    label=semana,
                    command=lambda v=semana: actualizar_seleccion("Semana", v)
                )

    # Detectar y normalizar valores Ãºnicos de otras columnas
    columnas_procesar = [
        "TIPO SICE", "ESTADO SICE", "TIPO MMS", "Estado MMS",
        "SISTEMA", "TIPO", "NOMBRE DEL EQUIPO",
        "OTROS SISTEMAS", "EMPLAZAMIENTO", "LINEA", "TRATAMIENTO"
    ]

    for columna in columnas_procesar:
        if columna in data.columns:
            valores_originales = data[columna].dropna().unique().tolist()
            normalizados = {normalizar_texto(valor): valor for valor in valores_originales}
            valores_unicos = list(normalizados.values())
            valores_unicos.sort()
            valores_unicos.insert(0, "TODOS")  # Agregar opciÃ³n "TODOS"
            

            if columna not in variables:
                variables[columna] = tk.StringVar(value="Seleccione")

            if columna not in menus:
                label = tk.Label(frame_campos, text=columna, bg="#2b2b2b", fg="white", font=("arial", 10))
                label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
                menus[columna] = tk.OptionMenu(frame_campos, variables[columna], *valores_unicos)
                menus[columna].grid(row=2, column=1, padx=10, pady=5, sticky="w")
            else:
                menus[columna].children["menu"].delete(0, "end")
                for valor in valores_unicos:
                    menus[columna].children["menu"].add_command(
                        label=valor,
                        command=lambda v=valor, col=columna: actualizar_seleccion(col, v)
                    )


# FunciÃ³n para actualizar las selecciones y previsualizarlas
def actualizar_seleccion(campo, valor):
    variables[campo].set(valor)

# Modificar la funciÃ³n exportar_seleccion
def exportar_seleccion():
    if data is None or data.empty:
        messagebox.showwarning("Advertencia", "No hay datos cargados para filtrar.")
        return

    # Recoger las selecciones de los campos generales
    seleccion = {campo: variables[campo].get() for campo in variables if variables[campo].get() != "Seleccione"}

    # Recoger palabras clave de los campos de descripciÃ³n
    palabras_clave_descripcion = {campo: variables_independientes[campo].get().strip().lower()
                                   for campo in campos_independientes if variables_independientes[campo].get().strip()}

    # Crear una copia del DataFrame original
    datos_filtrados = data.copy()

    # Normalizar las columnas relevantes
    for columna in datos_filtrados.columns:
        if datos_filtrados[columna].dtype == object:  # Normalizar solo columnas de tipo texto
            datos_filtrados[columna] = datos_filtrados[columna].fillna("").apply(normalizar_texto)

    # Aplicar los filtros por selecciones generales
    for columna, valor in seleccion.items():
        if columna in datos_filtrados.columns and valor != "TODOS":
            if columna == "Fecha" and valor != "TODOS":  # Filtro especial para el aÃ±o
                try:
                    valor = int(valor)  # Convertir el aÃ±o seleccionado a entero
                    fecha_inicio = f"{valor}-01-01"  # Inicio del aÃ±o
                    fecha_fin = f"{valor}-12-31"  # Fin del aÃ±o

                    # Filtrar datos dentro del rango del aÃ±o
                    datos_filtrados = datos_filtrados[
                        (datos_filtrados["Fecha"] >= pd.to_datetime(fecha_inicio)) &
                        (datos_filtrados["Fecha"] <= pd.to_datetime(fecha_fin))
                    ]
                except ValueError:
                    messagebox.showerror("Error", "El valor del aÃ±o seleccionado no es vÃ¡lido.")
                    return
            else:  # Filtro estÃ¡ndar para otras columnas
                datos_filtrados = datos_filtrados[
                    datos_filtrados[columna] == normalizar_texto(valor)
                ]

    # Aplicar filtros basados en palabras clave de los campos de descripciÃ³n
    for columna, palabra_clave in palabras_clave_descripcion.items():
        if columna in datos_filtrados.columns:  # Verificar que la columna exista
            if palabra_clave:
                # Dividir las palabras clave ingresadas por el usuario en una lista
                palabras = [normalizar_texto(p) for p in palabra_clave.strip().split()]

                # LÃ³gica de filtro: Y u O
                if logica_filtro.get() == "Y":  # LÃ³gica AND (todas las palabras deben coincidir)
                    for palabra in palabras:
                        datos_filtrados = datos_filtrados[
                            datos_filtrados[columna].str.contains(palabra, na=False)
                        ]
                elif logica_filtro.get() == "O":  # LÃ³gica OR (al menos una palabra debe coincidir)
                    datos_filtrados = datos_filtrados[
                        datos_filtrados[columna].apply(lambda x: any(palabra in x for palabra in palabras))
                    ]

    # Verificar si hay datos despuÃ©s de aplicar los filtros
    if datos_filtrados.empty:
        messagebox.showinfo("Sin resultados", "No hay datos que coincidan con los filtros seleccionados.")
        return

    # Cambiar el nombre de 'Unnamed: 0' a 'NÂ°' y enumerar las filas
    if "Unnamed: 0" in datos_filtrados.columns:
        datos_filtrados = datos_filtrados.rename(columns={"Unnamed: 0": "NÂ°"})
    datos_filtrados["NÂ°"] = range(1, len(datos_filtrados) + 1)  # Enumerar las filas

    # Asegurarse de que la columna Fecha solo muestre la fecha (sin hora)
    if "Fecha" in datos_filtrados.columns:
        datos_filtrados["Fecha"] = datos_filtrados["Fecha"].dt.date  # Extraer solo la fecha

    # Guardar los datos filtrados en un archivo Excel
    save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
    if save_path:
        datos_filtrados.to_excel(save_path, index=False)  # Exportar usando pandas

        # Ajustar el ancho de columnas con openpyxl
        try:
            wb = load_workbook(save_path)  # Cargar el archivo guardado
            ws = wb.active  # Obtener la hoja activa

            # Ajustar el ancho de cada columna segÃºn su contenido
            for column_cells in ws.columns:
                max_length = 0
                column_letter = column_cells[0].column_letter  # Letra de la columna (A, B, C, ...)
                for cell in column_cells:
                    if cell.value:  # Si la celda tiene un valor
                        max_length = max(max_length, len(str(cell.value)))
                adjusted_width = max_length + 2  # Ajustar el ancho con un pequeÃ±o margen
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(save_path)  # Guardar los cambios
            wb.close()
            messagebox.showinfo("Ã‰xito", f"Archivo exportado correctamente en:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo ajustar el tamaÃ±o de las columnas: {e}")

# ConfiguraciÃ³n de la ventana principal
ventana = tk.Tk()
ventana.title("SICE INFORMES 2025")
ventana.geometry("1000x600")
ventana.configure(bg="#2b2b2b")

# Frame del menÃº principal
frame_menu = tk.Frame(ventana, bg="#2b2b2b")
label_menu = tk.Label(frame_menu, text="SICE 2025\nMenÃº Principal", bg="#2b2b2b", fg="white", font=("arial", 20, "bold"))
label_menu.pack(pady=20)

boton_anexo = tk.Button(
    frame_menu, 
    text="Ir a Informe anexos", 
    command=lambda: [frame_menu.pack_forget(), frame_anexo.pack(fill="both", expand=True)],
    font=("arial", 10, "bold")
)
boton_anexo.pack(pady=10)

boton_nuevo = tk.Button(
    frame_menu, 
    text="Informe semanal", 
    command=frame_semanal,
    font=("arial", 10, "bold")
)
boton_nuevo.pack(pady=10)

# Frame para la opciÃ³n "Informe semanal"
frame_anexo = tk.Frame(ventana, bg="#2b2b2b")
label_anexo = tk.Label(
    frame_anexo,
    text="GestiÃ³n de Informes Semanales",
    bg="#2b2b2b",
    fg="white",
    font=("arial", 20, "bold")
)
label_anexo.pack(pady=20)


# Frame para la opciÃ³n "Informe semanal"
frame_anexo = tk.Frame(ventana, bg="#2b2b2b")

# ConfiguraciÃ³n para centrar todo el contenido en el frame_anexo
frame_anexo.grid_rowconfigure(0, weight=1)  # Espacio arriba
frame_anexo.grid_rowconfigure(1, weight=1)  # TÃ­tulo
frame_anexo.grid_rowconfigure(2, weight=1)  # BotÃ³n insertar
frame_anexo.grid_rowconfigure(3, weight=2)  # Campos principales
frame_anexo.grid_rowconfigure(4, weight=1)  # BotÃ³n exportar
frame_anexo.grid_rowconfigure(5, weight=1)  # BotÃ³n volver
frame_anexo.grid_columnconfigure(0, weight=1)  # Centrar horizontalmente

# TÃ­tulo
label_anexo = tk.Label(
    frame_anexo,
    text="GestiÃ³n de Informes Semanales",
    bg="#2b2b2b",
    fg="white",
    font=("arial", 20, "bold")
)
label_anexo.grid(row=1, column=0, pady=10, sticky="n")

# Etiqueta de informaciÃ³n adicional
label_informacion = tk.Label(
    frame_anexo,
    text="Rellene los campos correspondientes.",
    bg="#1e1e1e",
    fg="#a9a9a9",
    font=("arial", 10, "italic")
)
label_informacion.grid(row=2, column=0, pady=5, sticky="n")

# BotÃ³n para cargar archivo Excel
boton_insertar = tk.Button(
    frame_anexo,
    text="Cargar Archivo Excel",
    command=archivo_anexo,
    font=("arial", 10, "bold"),
    activebackground="#4c70ba",
    activeforeground="white",
    bd=2,
    relief="raised"
)
boton_insertar.grid(row=3, column=0, pady=15, sticky="n")

# Frame para organizar las entradas en columnas
frame_campos = tk.Frame(frame_anexo, bg="#2b2b2b")
frame_campos.grid(row=4, column=0, pady=10, sticky="nsew")

# ConfiguraciÃ³n del frame_campos para centrar sus elementos
frame_campos.grid_rowconfigure(0, weight=0)
frame_campos.grid_columnconfigure((0, 1, 2, 3, 4), weight=0)  # Espaciado uniforme

# Variables para los campos generales
campos = [
    "Fecha", "Semana", "TIPO SICE", "ESTADO SICE", "TIPO MMS", "Estado MMS",
    "SISTEMA", "TIPO", "NOMBRE DEL EQUIPO", "OTROS SISTEMAS", 
    "EMPLAZAMIENTO", "LINEA", "TRATAMIENTO"
]
variables = {campo: tk.StringVar(value="Seleccione") for campo in campos}
menus = {}

# Crear las columnas principales y centrarlas
num_columnas = 3
total_columnas = num_columnas * 2 + 2  # AÃ±adir dos columnas vacÃ­as (izquierda y derecha)

# Configurar las columnas para centrar el contenido
for col in range(total_columnas):
    if col == 0 or col == total_columnas - 1:  # Columnas vacÃ­as (izquierda y derecha)
        frame_campos.grid_columnconfigure(col, weight=1)  # Espacio expansible
    else:  # Columnas con contenido
        frame_campos.grid_columnconfigure(col, weight=0)  # Fijas

# Crear las columnas principales centradas
for i, campo in enumerate(campos):
    columna = (i % num_columnas) * 2 + 1  # Ajustar posiciÃ³n dejando espacio a los lados
    fila = i // num_columnas

    # Crear etiquetas alineadas a la derecha
    label = tk.Label(
        frame_campos,
        text=campo,
        bg="#2b2b2b",
        fg="white",
        font=("arial", 10)
    )
    label.grid(row=fila, column=columna, padx=10, pady=5, sticky="e")  # Alinear a la derecha

    # Crear menÃºs desplegables alineados a la izquierda
    menu = tk.OptionMenu(frame_campos, variables[campo], "Seleccione", command=lambda v, c=campo: actualizar_seleccion(c, v))
    menu.config(width=15, bg="white", fg="#2b2b2b")
    menu.grid(row=fila, column=columna + 1, padx=10, pady=5, sticky="w")  # Alinear a la izquierda
    menus[campo] = menu

# Crear un marco adicional para los campos de descripciÃ³n independientes
frame_independiente = tk.Frame(frame_campos, bg="#2b2b2b")
frame_independiente.grid(row=fila + 1, column=0, columnspan=num_columnas * 2, pady=20, sticky="nsew")

# Variables para los campos independientes
campos_independientes = ["DESCRIPCION", "DESCRIPCIÃ“N DE LA FALLA"]
variables_independientes = {campo: tk.StringVar(value="") for campo in campos_independientes}

# ConfiguraciÃ³n del frame_independiente para centrar los campos de descripciÃ³n
total_columnas_descripcion = 4  # Dos columnas vacÃ­as a los lados y dos para contenido

# Configurar las columnas para centrar el contenido
for col in range(total_columnas_descripcion):
    if col == 0 or col == total_columnas_descripcion - 1:  # Columnas vacÃ­as (izquierda y derecha)
        frame_independiente.grid_columnconfigure(col, weight=1)  # Espacio expansible
    else:  # Columnas con contenido
        frame_independiente.grid_columnconfigure(col, weight=0)  # Fijas

# Centrando los campos de descripciÃ³n
for i, campo in enumerate(campos_independientes):
    # Etiquetas alineadas a la derecha
    label = tk.Label(
        frame_independiente,
        text=campo,
        bg="#2b2b2b",
        fg="white",
        font=("arial", 10)
    )
    label.grid(row=i, column=1, padx=10, pady=5, sticky="e")  # Columna 1: etiquetas

    # Entradas alineadas a la izquierda
    entry = tk.Entry(
        frame_independiente,
        textvariable=variables_independientes[campo],
        width=40,
        bg="white",
        fg="#2b2b2b",
        font=("arial", 10)
    )
    entry.grid(row=i, column=2, padx=10, pady=5, sticky="w")  # Columna 2: entradas

def actualizar_logica(valor):
    logica_filtro.set(valor)

# Variables globales
df_averias = None
df_filtrado = None
df_tabla12prox = None
df_averias_mes = None

def abrir_averias():
    global df_averias, df_averias_mes
    ruta_archivo = filedialog.askopenfilename(filetypes=[("Archivos Excel", "*.xlsx")])
    if not ruta_archivo:
        return  
    df_otro = pd.read_excel(ruta_archivo, dtype=str, skiprows=4, usecols="B:Y")
    columnas_necesarias = ["LINEA", "EMPLAZAMIENTO", "OT", "DESCRIPCIÃ“N DE LA FALLA", "ACTIVO", "CAT ", "TIPO", "FECHA HORA INFORME", "ESTADO SICE", "Semana", "TIPO SICE", "SISTEMA"]
    if not all(col in df_otro.columns for col in columnas_necesarias):
        columnas_faltantes = [col for col in columnas_necesarias if col not in df_otro.columns]
        messagebox.showerror("Error", f"Faltan las siguientes columnas en el archivo: {', '.join(columnas_faltantes)}")
        return
    df_otro["FECHA HORA INFORME"] = pd.to_datetime(df_otro["FECHA HORA INFORME"], errors="coerce")
    if df_otro["FECHA HORA INFORME"].isnull().all():
        messagebox.showerror("Error", "La columna 'FECHA HORA INFORME' no contiene fechas vÃ¡lidas.")
        return
    fecha_minima = df_otro["FECHA HORA INFORME"].min()
    fecha_maxima = df_otro["FECHA HORA INFORME"].max()
    ventana_fechas = tk.Toplevel()
    ventana_fechas.title("Filtrar por Fecha de Informe")
    ventana_fechas.geometry("800x700")
    tk.Label(ventana_fechas, text="Fecha de inicio:").pack(pady=5)
    calendario_inicio = Calendar(ventana_fechas, selectmode="day", date_pattern="yyyy-mm-dd", mindate=fecha_minima, maxdate=fecha_maxima)
    calendario_inicio.pack(pady=10)
    tk.Label(ventana_fechas, text="Fecha de fin:").pack(pady=5)
    calendario_fin = Calendar(ventana_fechas, selectmode="day", date_pattern="yyyy-mm-dd", mindate=fecha_minima, maxdate=fecha_maxima)
    calendario_fin.pack(pady=10)
    def aplicar_filtro():
        global df_averias, df_averias_mes  # Se agregarÃ¡ df_averias_mes para el segundo filtro
        ventana_fechas.destroy()
        fecha_inicio = datetime.strptime(calendario_inicio.get_date(), "%Y-%m-%d")
        fecha_fin = datetime.strptime(calendario_fin.get_date(), "%Y-%m-%d") + timedelta(days=1) - timedelta(seconds=1)
        if fecha_inicio > fecha_fin:
            messagebox.showerror("Error", "La fecha de inicio no puede ser mayor que la fecha de fin.")
            return
        # ðŸ”¹ Filtrar por rango de fechas seleccionado por el usuario
        df_filtrado = df_otro[(df_otro["FECHA HORA INFORME"] >= fecha_inicio) & 
                              (df_otro["FECHA HORA INFORME"] <= fecha_fin)]
        if df_filtrado.empty:
            messagebox.showinfo("Sin datos", "No se encontraron averÃ­as en el rango de fechas seleccionado.")
            return
        df_averias = df_filtrado[columnas_necesarias]
        messagebox.showinfo("Ã‰xito", "Los datos se han filtrado correctamente.")
        # ðŸ”¹ Ahora aplicar el filtro para obtener el Ãºltimo mes completo con datos todos los dÃ­as
        obtener_mes_con_datos(df_otro)
    tk.Button(ventana_fechas, text="Aplicar Filtro", command=aplicar_filtro).pack(pady=10)

    def obtener_mes_con_datos(df):
        global df_averias_mes
        """
        Filtra el DataFrame para obtener los datos del mes correspondiente a la fecha mÃ¡xima.
        Si el Ãºltimo dÃ­a del mes no tiene datos, se selecciona el mes anterior.
        """
        # Convertir la columna de fecha a tipo datetime si no lo estÃ¡
        df["FECHA HORA INFORME"] = pd.to_datetime(df["FECHA HORA INFORME"], errors="coerce")

        # Obtener la fecha mÃ¡xima del DataFrame
        fecha_maxima = df["FECHA HORA INFORME"].max()

        if pd.isna(fecha_maxima):
            print("No hay fechas disponibles en los datos.")
            return pd.DataFrame()  # Retornar DataFrame vacÃ­o

        # Obtener el Ãºltimo dÃ­a del mes de la fecha mÃ¡xima
        ultimo_dia_mes = pd.Timestamp(fecha_maxima.year, fecha_maxima.month, 1).days_in_month
        fecha_ultimo_dia = pd.Timestamp(fecha_maxima.year, fecha_maxima.month, ultimo_dia_mes)

        # Verificar si hay datos en el Ãºltimo dÃ­a del mes
        tiene_datos_ultimo_dia = not df[df["FECHA HORA INFORME"].dt.date == fecha_ultimo_dia.date()].empty

        if tiene_datos_ultimo_dia:
            # âœ… Si el Ãºltimo dÃ­a tiene datos, filtrar SOLO el mes de la fecha mÃ¡xima
            df_averias_mes = df[
                (df["FECHA HORA INFORME"].dt.year == fecha_maxima.year) &
                (df["FECHA HORA INFORME"].dt.month == fecha_maxima.month)
            ]
            print(f"Se encontraron datos para {fecha_maxima.month}/{fecha_maxima.year}.")
        else:
            # âŒ Si NO hay datos el Ãºltimo dÃ­a del mes, tomar el mes anterior
            mes_anterior = fecha_maxima.replace(day=1) - timedelta(days=1)  # Ãšltimo dÃ­a del mes anterior
            df_averias_mes = df[
                (df["FECHA HORA INFORME"].dt.year == mes_anterior.year) &
                (df["FECHA HORA INFORME"].dt.month == mes_anterior.month)]
            print(f"Se encontraron datos para el mes anterior {mes_anterior.month}/{mes_anterior.year}.")

def abrir_programacion():
    global df_filtrado, df_tabla12prox  
    try:
        ruta_archivo = filedialog.askopenfilename(filetypes=[("Archivos Excel", "*.xlsx")])
        if not ruta_archivo:
            return

        df = pd.read_excel(ruta_archivo, sheet_name="L63")
        df.columns = df.iloc[0]
        df = df[1:]

        if "FE" not in df.columns:
            messagebox.showerror("Error", "El archivo no contiene una columna 'FE'.")
            return

        df["FE"] = pd.to_datetime(df["FE"], errors="coerce")

        if df["FE"].isnull().all():
            messagebox.showerror("Error", "La columna 'FE' no contiene fechas vÃ¡lidas.")
            return

        fecha_minima = df["FE"].min()
        fecha_maxima = df["FE"].max()

        if pd.isnull(fecha_minima) or pd.isnull(fecha_maxima):
            messagebox.showerror("Error", "No se encontraron fechas vÃ¡lidas para establecer el rango.")
            return

        ventana_fechas = tk.Toplevel()
        ventana_fechas.title("Filtrar por Fecha")
        ventana_fechas.geometry("800x700")

        tk.Label(ventana_fechas, text="Fecha de inicio:").pack(pady=5)
        calendario_inicio = Calendar(
            ventana_fechas, selectmode="day", date_pattern="yyyy-mm-dd",
            mindate=min(fecha_minima, datetime.now() - pd.Timedelta(days=365)),
            maxdate=max(fecha_maxima, datetime.now() + pd.Timedelta(days=365))
        )
        calendario_inicio.pack(pady=10)

        tk.Label(ventana_fechas, text="Fecha de fin:").pack(pady=5)
        calendario_fin = Calendar(
            ventana_fechas, selectmode="day", date_pattern="yyyy-mm-dd",
            mindate=min(fecha_minima, datetime.now() - pd.Timedelta(days=365)),
            maxdate=max(fecha_maxima, datetime.now() + pd.Timedelta(days=365))
        )
        calendario_fin.pack(pady=10)

        def aplicar_filtro():
            global df_filtrado, df_tabla12prox  
            ventana_fechas.destroy()
            try:
                fecha_inicio = datetime.strptime(calendario_inicio.get_date(), "%Y-%m-%d")
                fecha_fin = datetime.strptime(calendario_fin.get_date(), "%Y-%m-%d")

                if fecha_inicio > fecha_fin:
                    messagebox.showerror("Error", "La fecha de inicio no puede ser mayor que la fecha de fin.")
                    return

                df_filtrado = df[(df["FE"] >= fecha_inicio) & (df["FE"] <= fecha_fin)]

                if df_filtrado.empty:
                    messagebox.showinfo("Sin datos", "No se encontraron datos para el rango de fechas seleccionado.")
                    return

                df_filtrado = df_filtrado[["Descripcion OT", "Equipo", "DEP", "EST", "SIST", "F.LIBERACIÃ“N", "FP", "NÃºmero de OT", "FE", "CAT", "ESTADO MMS"]]
                df_filtrado["FE"] = df_filtrado["FE"].dt.strftime('%Y-%m-%d')

            except ValueError:
                messagebox.showerror("Error", "Error al procesar las fechas seleccionadas.")

            df_tabla12prox = pd.read_excel(ruta_archivo, sheet_name="INFORME-PRÃ“XIMAS 12SEM", skiprows=6)

            df_tabla12prox = df_tabla12prox.iloc[:, 11:].reset_index(drop=True)

            columna_titulo = df_tabla12prox.columns[0]
            fila_inicio = df_tabla12prox[df_tabla12prox[columna_titulo] == "ProyecciÃ³n de actividades segÃºn su categorÃ­a"].index

            if not fila_inicio.empty:
                fila_inicio = fila_inicio[0] + 1
                df_tabla12prox = df_tabla12prox.iloc[fila_inicio:].reset_index(drop=True)
                df_tabla12prox.dropna(axis=1, how='all', inplace=True)
                df_tabla12prox.columns = df_tabla12prox.iloc[0].fillna("")
                df_tabla12prox = df_tabla12prox[1:].reset_index(drop=True)
                df_tabla12prox = df_tabla12prox.loc[:, ~df_tabla12prox.columns.astype(str).str.contains("Unnamed|^$", regex=True)]

            messagebox.showinfo("Ã‰xito", "ProgramaciÃ³n cargada correctamente.")

        boton_aplicar = tk.Button(ventana_fechas, text="Aplicar Filtro", command=aplicar_filtro)
        boton_aplicar.pack(pady=10)

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo abrir el archivo Excel:\n{e}")


def crear_word():
    global df_tabla12prox
    # Crear un nuevo documento de Word
    doc = Document()

    doc.add_paragraph('\n'*2)

    def semanas_transcurridas():
        # Obtener el aÃ±o actual
        aÃ±o = datetime.now().year
        # Crear el objeto de fecha para el primer dÃ­a del aÃ±o
        fecha_inicial_obj = datetime.strptime(str(aÃ±o)+("/01/01"), "%Y/%m/%d")

        # Obtener la fecha actual
        fecha_final_obj = datetime.now()  

        # Calcular la diferencia en dÃ­as entre la fecha final y la inicial
        diferencia_dias = (fecha_final_obj - fecha_inicial_obj).days + 1  # Se suma 1 para contar el primer dÃ­a como parte de la semana

        # Calcular las semanas transcurridas (redondeando hacia arriba si es necesario)
        semanas_transcurridas = diferencia_dias // 7
        if diferencia_dias % 7 != 0:
            semanas_transcurridas += 1  # Si no es exacto, contar la semana adicional
        return str(semanas_transcurridas)

    # Para centrar la imagen, puedes agregar un parÃ¡grafo vacÃ­o antes y despuÃ©s de la imagen
    # y luego utilizar el mÃ©todo `alignment` para centrar el parÃ¡grafo
    imagen_portada = doc.add_paragraph()
    imagen_portada.alignment = 1  # 1 es el valor para centrar
    imagen_portada.add_run().add_picture("imagenes/logo_metro.png", width=Inches(3))
    imagen_portada = doc.add_paragraph()

    portada = doc.add_paragraph('\nMANTENIMIENTO DEL SISTEMA DE COMUNICACIONES \nLINEAS 6 Y 3\nMETRO DE SANTIAGO \n\nREPORTE SEMANAL ' + str(datetime.now().year) + '\nSEMANA ' + semanas_transcurridas() + '\nCONTRATO NÂ° MN-236-2014-G')
    portada.alignment = 1
    portada.runs[0].font.name = 'Calibri'
    portada.runs[0].font.size = Pt(15.5)
    portada.runs[0].font.bold = True

    # Acceder al pie de pÃ¡gina de la primera secciÃ³n
    footer = doc.sections[0].footer

    # AÃ±adir un pÃ¡rrafo al pie de pÃ¡gina
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  

    # Crear un campo para el nÃºmero de pÃ¡gina
    page_field = OxmlElement('w:fldSimple')
    page_field.set(qn('w:instr'), 'PAGE')

    # AÃ±adir el campo al pÃ¡rrafo
    run = p.add_run()
    run._r.append(page_field)

    # Opcional: Personalizar el estilo del texto
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)

    # Acceder a la secciÃ³n del encabezado
    section = doc.sections[0]
    header = section.header

    # Crear una tabla en el encabezado con 3 celdas (izquierda, centro, derecha)
    table = header.add_table(rows=1, cols=3, width=Inches(6))
    table.autofit = False

    # Configurar el ancho de las columnas
    for idx, width in enumerate([Inches(2), Inches(2), Inches(2)]):
        table.columns[idx].width = width

    # Obtener la primera fila de la tabla
    row = table.rows[0].cells

    # Insertar la imagen en la celda izquierda
    celda_izquierda = row[0].paragraphs[0]
    try:
        celda_izquierda.add_run().add_picture("imagenes/logo_metro.png", width=Inches(1.2))
        celda_izquierda.paragraph_format.alignment = 0
    except Exception as e:
        print(f"Error al cargar la imagen izquierda: {e}")

    # Insertar el texto en la celda central
    celda_centro = row[1].paragraphs[0]
    celda_centro.text = "METRO DE SANTIAGO LÃNEAS 6 Y 3 \nINFORME SEMANA " + semanas_transcurridas() + "\n Contrato NÂ° MN-236-2014-G" 
    celda_centro.runs[0].font.name = 'Calibri'
    celda_centro.runs[0].font.size = Pt(8.5)
    celda_centro.paragraph_format.alignment = 1  # Centrar el texto (1 = CENTER)

    # Insertar la imagen en la celda derecha
    celda_derecha = row[2].paragraphs[0]
    try:
        celda_derecha.add_run().add_picture("imagenes/logo_sice.png", width=Inches(1))
        celda_derecha.paragraph_format.alignment = 2
    except Exception as e:
        print(f"Error al cargar la imagen derecha: {e}")
    
    doc.add_paragraph("\n" * 8)
    # AÃ±adir una tabla con el nÃºmero de filas y columnas necesarias
    tabla_portada = doc.add_table(rows=4, cols=7)
    tabla_portada.style = 'Table Grid'

    # Modificar el ancho de las celdas (opcional)
    for row in tabla_portada.rows:
        for cell in row.cells:
            cell.width = Pt(70)

    # Contenido de las celdas
    header_row = ['\n'+str(Rev_documento.get())+'\n', '09-12-2024', 'REVISIÃ“N', 'HSANMARTIN', 'RZAMBRANO', 'MCOLLAO', 'LLÃ“PEZ']
    sub_header_row = ['REV.', 'FECHA', 'EMITIDO PARA', 'PREPARÃ“', 'REVISÃ“', 'APROBÃ“', 'APROBÃ“ METRO']
    footer_row = ['NÂº DOCUMENTO', '', 'P63-MA-0632-INT-000-CO-00347', '', '\nPÃ¡gina 1 de 20', '', 'REV. ' + str(Rev_documento.get())]

    # FunciÃ³n para dar formato a las celdas con estilo de letra
    def format_cell(cell, text, font_name="Calibri", font_size=9.5, font_bold= False):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font_name  # Cambiar la fuente
        run.font.size = Pt(font_size)  # Cambiar el tamaÃ±o de letra
        run.font.bold = font_bold  # Negrita
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # AÃ±adir datos y aplicar estilo a cada celda
    for i, text in enumerate(header_row):
        format_cell(tabla_portada.rows[0].cells[i], text, font_name="Calibri", font_size=9.5, font_bold=False)
    for i, text in enumerate(sub_header_row):
        format_cell(tabla_portada.rows[1].cells[i], text, font_name="Calibri", font_size=9.5, font_bold =True)
    for i, text in enumerate(footer_row):
        if i == 0 or i == 6:
            format_cell(tabla_portada.rows[2].cells[i], text, font_name="Calibri", font_size=9.5, font_bold=True)
        else:
            format_cell(tabla_portada.rows[2].cells[i], text, font_name="Calibri", font_size=9.5, font_bold=False)

    logo_tabla = tabla_portada.rows[3].cells[6].paragraphs[0]
    logo_tabla.add_run().add_picture("imagenes/logo_sice.png", width=Inches(1))
    logo_tabla.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Unir celdas 
    tabla_portada.rows[2].cells[4].merge(tabla_portada.rows[2].cells[5])
    tabla_portada.rows[2].cells[1].merge(tabla_portada.rows[2].cells[3])
    tabla_portada.rows[3].cells[0].merge(tabla_portada.rows[3].cells[6])
    
    doc.add_page_break()   

    # PAGINA 2

    Control_cambios = doc.add_paragraph('\nCONTROL DE CAMBIOS')
    Control_cambios.alignment = 1
    Control_cambios.runs[0].font.name = 'Calibri'
    Control_cambios.runs[0].font.size = Pt(10)
    Control_cambios.runs[0].font.bold = True

    tabla_2 = doc.add_table(rows=2, cols=7)
    tabla_2.style = 'Table Grid'

    Rev_titulo = ['REVISIÃ“N', 'Fecha', 'PÃGINA', 'PÃRRAFO', '', 'MODIFICACIÃ“N REALIZADA', '']
    Rev_contenido= [str(Rev_documento.get()), str(datetime.today().date()), 'N/A', 'N/A', '', 'PRIMERA VERSIÃ“N DEL DOCUMENTO', '']

    for i, text in enumerate(Rev_titulo):
        format_cell(tabla_2.rows[0].cells[i], text, font_name="Calibri", font_size=9.5, font_bold=False)
    for i, text in enumerate(Rev_contenido):
        format_cell(tabla_2.rows[1].cells[i], text, font_name="Calibri", font_size=9.5, font_bold =False)
    tabla_2.rows[0].cells[4].merge(tabla_portada.rows[0].cells[6])
    tabla_2.rows[1].cells[4].merge(tabla_portada.rows[1].cells[6])
    doc.add_page_break()
    # indice

    titulo_indice = doc.add_heading('Indice')
    titulo_indice.runs[0].font.name = 'Calibri'
    titulo_indice.runs[0].font.size = Pt(10)
    titulo_indice.runs[0].font.bold = True
    titulo_indice.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro
    titulo_indice.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()    

    # Objetivo
    titulo_objetivo = doc.add_heading('1. Objetivo', 1)
    titulo_objetivo.runs[0].font.name = 'Calibri'
    titulo_objetivo.runs[0].font.size = Pt(10)
    titulo_objetivo.runs[0].font.bold = True
    titulo_objetivo.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro
    objeto = doc.add_paragraph('\n        Reportar el desarrollo de las actividades de Mantenimiento Preventivo, Correctivo, eventos y Hallazgos asociados a la gestiÃ³n del contrato de Mantenimiento MN-236-2014-G correspondiente a la Semana '+ semanas_transcurridas()+ ' del aÃ±o ' +str(datetime.now().year)+ ' del sistema de Comunicaciones de las lÃ­neas 6 y 3 del Metro de Santiago.')
    objeto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    objeto.runs[0].font.name = 'Calibri'
    objeto.runs[0].font.size = Pt(9.5)
    objeto.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    doc.add_page_break()   

    # Mantenimiento
    titulo_mantenimiento = doc.add_heading('2. Mantenimiento Preventivo', 1)
    titulo_mantenimiento.runs[0].font.name = 'Calibri'
    titulo_mantenimiento.runs[0].font.size = Pt(10)
    titulo_mantenimiento.runs[0].font.bold = True
    titulo_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    titulo_mantenimiento2_1 = doc.add_heading('2.1.   ProgramaciÃ³n Semanal del Mantenimiento Preventivo ', 2)
    titulo_mantenimiento2_1.runs[0].font.name = 'Calibri'
    titulo_mantenimiento2_1.runs[0].font.size = Pt(10)
    titulo_mantenimiento2_1.runs[0].font.bold = True
    titulo_mantenimiento2_1.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento2_1 = doc.add_paragraph('A continuaciÃ³n, se presentan los mantenimientos preventivos segÃºn su categorÃ­a que fueron planificados para la Semana ' + semanas_transcurridas()+ ' del aÃ±o ' +str(datetime.now().year))
    mantenimiento2_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mantenimiento2_1.runs[0].font.name = 'Calibri'
    mantenimiento2_1.runs[0].font.size = Pt(9.5)

    titulo_mantenimiento_CAT1 = doc.add_paragraph('\nMantenimientos Preventivos CAT 1')
    titulo_mantenimiento_CAT1.runs[0].font.name = 'Calibri'
    titulo_mantenimiento_CAT1.runs[0].font.size = Pt(10)
    titulo_mantenimiento_CAT1.runs[0].font.bold = True
    titulo_mantenimiento_CAT1.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro
    titulo_mantenimiento_CAT1.alignment = 1

    mantenimiento_CAT1 = doc.add_paragraph('\n      Para la semana en estudio se registran los siguientes mantenimientos preventivos de categorÃ­a 1.')
    mantenimiento_CAT1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mantenimiento_CAT1.runs[0].font.name = 'Calibri'
    mantenimiento_CAT1.runs[0].font.size = Pt(9.5)

#
#
    try:
        if not df_filtrado.empty:
            # Filtrar los valores de CAT1
            if 'CAT' in df_filtrado.columns:
                df_filtro1 = df_filtrado[df_filtrado['CAT'] == 'CAT 1']

            # Eliminar las columnas no requeridas del DataFrame si existen
            columnas_requeridas = ['Descripcion OT', 'Equipo', 'DEP', 'EST', 'SIST', 'FP', 'FE', 'NÃºmero de OT']
            columnas_a_eliminar = [col for col in df_filtro1.columns if col not in columnas_requeridas]
            df_filtro1 = df_filtro1.drop(columns=columnas_a_eliminar, errors='ignore')

            # Formatear las columnas 'FE' y 'FP' para que solo contengan la fecha
            for col in ['FE', 'FP']:
                if col in df_filtro1.columns:
                    df_filtro1[col] = pd.to_datetime(df_filtro1[col], errors='coerce').dt.date

            # Crear la tabla con las columnas requeridas
            columnas = ["NÂ°", "DESCRIPCIÃ“N", "Equipo", "LÃ­nea", "EstaciÃ³n", "Sistema", "Plan Matriz", "Ejecutado", "OT"]
            num_filas = df_filtro1.shape[0]  # NÃºmero de filas de datos
            num_columnas = len(columnas)  # NÃºmero de columnas predefinidas

            # Ordenar las columnas segÃºn el formato deseado
            df_filtro1 = df_filtro1[['Descripcion OT', 'Equipo', 'DEP', 'EST', 'SIST', 'FP', 'FE', 'NÃºmero de OT']]

            # Crear la tabla con las columnas requeridas
            columnas = ["NÂ°", "DESCRIPCIÃ“N", "Equipo", "LÃ­nea", "EstaciÃ³n", "Sistema", "Plan Matriz", "Ejecutado", "OT"]
            num_filas = df_filtro1.shape[0]  # NÃºmero de filas de datos
            num_columnas = len(columnas)  # NÃºmero de columnas predefinidas
            tablaCAT1 = doc.add_table(rows=num_filas + 1, cols=num_columnas)  # Crear tabla con encabezados y filas de datos
            tablaCAT1.style = 'Table Grid'

            # Agregar encabezados
            for j, column_name in enumerate(columnas):
                cell = tablaCAT1.cell(0, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Agregar los datos a la tabla
            for i, row in enumerate(df_filtro1.itertuples(index=False), start=1):
                # Agregar numeraciÃ³n en la primera columna
                cell = tablaCAT1.cell(i, 0)
                cell.text = str(i)
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Agregar el resto de los datos
                for j, value in enumerate(row, start=1):
                    cell = tablaCAT1.cell(i, j)
                    cell.text = str(value)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Ajustar altura de la fila para la columna de "DESCRIPCIÃ“N"
                    if columnas[j] == "DESCRIPCIÃ“N":
                        cell.width = Inches(3)  # Duplicar el espacio en la celda de descripciÃ³n

                    # Ajustar espacio para la columna "Plan Matriz"
                    if columnas[j] == "Plan Matriz":
                        cell.width = Inches(1.1)  

                # Centrar texto en todas las celdas
                for row in tablaCAT1.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')
    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")


    mantenimiento2_1 = doc.add_paragraph('Tabla NÂ°1 - Mantenimientos CategorÃ­a 1. \n')
    mantenimiento2_1.alignment = 1
    mantenimiento2_1.runs[0].font.name = 'Calibri'
    mantenimiento2_1.runs[0].font.size = Pt(9.5)

    titulo_mantenimiento_CAT2 = doc.add_paragraph('\nMantenimientos Preventivos CAT 2')
    titulo_mantenimiento_CAT2.runs[0].font.name = 'Calibri'
    titulo_mantenimiento_CAT2.runs[0].font.size = Pt(10)
    titulo_mantenimiento_CAT2.runs[0].font.bold = True
    titulo_mantenimiento_CAT2.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro
    titulo_mantenimiento_CAT2.alignment = 1

    mantenimiento_CAT2 = doc.add_paragraph('\n      Para la semana en estudio se registran los siguientes mantenimientos preventivos de categorÃ­a 2')
    mantenimiento_CAT2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mantenimiento_CAT2.runs[0].font.name = 'Calibri'
    mantenimiento_CAT2.runs[0].font.size = Pt(9.5)

    try:
        if not df_filtrado.empty:
            # Filtrar los valores de CAT1
            if 'CAT' in df_filtrado.columns:
                df_filtro1 = df_filtrado[df_filtrado['CAT'] == 'CAT 2']

            # Eliminar las columnas no requeridas del DataFrame si existen
            columnas_requeridas = ['Descripcion OT', 'Equipo', 'DEP', 'EST', 'SIST', 'FP', 'FE', 'NÃºmero de OT']
            columnas_a_eliminar = [col for col in df_filtro1.columns if col not in columnas_requeridas]
            df_filtro1 = df_filtro1.drop(columns=columnas_a_eliminar, errors='ignore')

            # Formatear las columnas 'FE' y 'FP' para que solo contengan la fecha
            for col in ['FE', 'FP']:
                if col in df_filtro1.columns:
                    df_filtro1[col] = pd.to_datetime(df_filtro1[col], errors='coerce').dt.date

            # Ordenar las columnas segÃºn el formato deseado
            df_filtro1 = df_filtro1[['Descripcion OT', 'Equipo', 'DEP', 'EST', 'SIST', 'FP', 'FE', 'NÃºmero de OT']]

            # Crear la tabla con las columnas requeridas
            columnas = ["NÂ°", "DESCRIPCIÃ“N", "Equipo", "LÃ­nea", "EstaciÃ³n", "Sistema", "Plan Matriz", "Ejecutado", "OT"]
            num_filas = df_filtro1.shape[0]  # NÃºmero de filas de datos
            num_columnas = len(columnas)  # NÃºmero de columnas predefinidas
            tablaCAT2 = doc.add_table(rows=num_filas + 1, cols=num_columnas)  # Crear tabla con encabezados y filas de datos
            tablaCAT2.style = 'Table Grid'

            # Agregar encabezados
            for j, column_name in enumerate(columnas):
                cell = tablaCAT2.cell(0, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Agregar los datos a la tabla
            for i, row in enumerate(df_filtro1.itertuples(index=False), start=1):
                # Agregar numeraciÃ³n en la primera columna
                cell = tablaCAT2.cell(i, 0)
                cell.text = str(i)
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Agregar el resto de los datos
                for j, value in enumerate(row, start=1):
                    cell = tablaCAT2.cell(i, j)
                    cell.text = str(value)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Ajustar altura de la fila para la columna de "DESCRIPCIÃ“N"
                    if columnas[j] == "DESCRIPCIÃ“N":
                        cell.width = Inches(3)  # Duplicar el espacio en la celda de descripciÃ³n
                    # Ajustar espacio para la columna "Plan Matriz"
                    if columnas[j] == "Plan Matriz":
                        cell.width = Inches(1.1)  

                # Centrar texto en todas las celdas
                for row in tablaCAT2.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')
    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")
    
    mantenimiento2_1 = doc.add_paragraph('Tabla NÂ°2 - Mantenimientos CategorÃ­a 2. \n')
    mantenimiento2_1.alignment = 1
    mantenimiento2_1.runs[0].font.name = 'Calibri'
    mantenimiento2_1.runs[0].font.size = Pt(9.5)

    titulo_mantenimiento_CAT3 = doc.add_paragraph('\nMantenimientos Preventivos CAT 3')
    titulo_mantenimiento_CAT3.runs[0].font.name = 'Calibri'
    titulo_mantenimiento_CAT3.runs[0].font.size = Pt(10)
    titulo_mantenimiento_CAT3.runs[0].font.bold = True
    titulo_mantenimiento_CAT3.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro
    titulo_mantenimiento_CAT3.alignment = 1

    mantenimiento_CAT2 = doc.add_paragraph('\n      Para la semana en estudio se registran los siguientes mantenimientos preventivos de categorÃ­a 3')
    mantenimiento_CAT2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mantenimiento_CAT2.runs[0].font.name = 'Calibri'
    mantenimiento_CAT2.runs[0].font.size = Pt(9.5)

    try:
        if not df_filtrado.empty:
            # Filtrar los valores de CAT1
            if 'CAT' in df_filtrado.columns:
                df_filtro1 = df_filtrado[df_filtrado['CAT'] == 'CAT 3']

            # Eliminar las columnas no requeridas del DataFrame si existen
            columnas_requeridas = ['Descripcion OT', 'Equipo', 'DEP', 'EST', 'SIST', 'FP', 'FE', 'NÃºmero de OT']
            columnas_a_eliminar = [col for col in df_filtro1.columns if col not in columnas_requeridas]
            df_filtro1 = df_filtro1.drop(columns=columnas_a_eliminar, errors='ignore')

            # Formatear las columnas 'FE' y 'FP' para que solo contengan la fecha
            for col in ['FE', 'FP']:
                if col in df_filtro1.columns:
                    df_filtro1[col] = pd.to_datetime(df_filtro1[col], errors='coerce').dt.date

            # Ordenar las columnas segÃºn el formato deseado
            df_filtro1 = df_filtro1[['Descripcion OT', 'Equipo', 'DEP', 'EST', 'SIST', 'FP', 'FE', 'NÃºmero de OT']]

            # Crear la tabla con las columnas requeridas
            columnas = ["NÂ°", "DESCRIPCIÃ“N", "Equipo", "LÃ­nea", "EstaciÃ³n", "Sistema", "Plan Matriz", "Ejecutado", "OT"]
            num_filas = df_filtro1.shape[0]  # NÃºmero de filas de datos
            num_columnas = len(columnas)  # NÃºmero de columnas predefinidas
            tablaCAT3 = doc.add_table(rows=num_filas + 1, cols=num_columnas)  # Crear tabla con encabezados y filas de datos
            tablaCAT3.style = 'Table Grid'

            # Agregar encabezados
            for j, column_name in enumerate(columnas):
                cell = tablaCAT3.cell(0, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Agregar los datos a la tabla
            for i, row in enumerate(df_filtro1.itertuples(index=False), start=1):
                # Agregar numeraciÃ³n en la primera columna
                cell = tablaCAT3.cell(i, 0)
                cell.text = str(i)
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Agregar el resto de los datos
                for j, value in enumerate(row, start=1):
                    cell = tablaCAT3.cell(i, j)
                    cell.text = str(value)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Ajustar altura de la fila para la columna de "DESCRIPCIÃ“N"
                    if columnas[j] == "DESCRIPCIÃ“N":
                        cell.width = Inches(3)  # Duplicar el espacio en la celda de descripciÃ³n

                    # Ajustar espacio para la columna "Plan Matriz"
                    if columnas[j] == "Plan Matriz":
                        cell.width = Inches(1.1)  

                # Centrar texto en todas las celdas
                for row in tablaCAT3.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')
    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")
    
    mantenimiento2_1 = doc.add_paragraph('Tabla NÂ°3 - Mantenimientos CategorÃ­a 3. \n')
    mantenimiento2_1.alignment = 1
    mantenimiento2_1.runs[0].font.name = 'Calibri'
    mantenimiento2_1.runs[0].font.size = Pt(9.5)

    doc.add_page_break()  

    #2.2 Adeherencia al plan de mantenimiento

    titulo_adherencia = doc.add_heading('2.2	 Adherencia Plan de Mantenimiento', 1)
    titulo_adherencia.runs[0].font.name = 'Calibri'
    titulo_adherencia.runs[0].font.size = Pt(10)
    titulo_adherencia.runs[0].font.bold = True
    titulo_adherencia.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    titulo_adherencia2_1 = doc.add_heading('2.2.1	Adherencia al Plan de Mantenimiento Semana ' + semanas_transcurridas(), 3)
    titulo_adherencia2_1.runs[0].font.name = 'Calibri'
    titulo_adherencia2_1.runs[0].font.size = Pt(10)
    titulo_adherencia2_1.runs[0].font.bold = True
    titulo_adherencia2_1.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    adherencia2_1 = doc.add_paragraph('     A continuaciÃ³n, se presenta tabla y grÃ¡fico donde se evidencia la adherencia al plan de Mantenimiento en la Semana ' + semanas_transcurridas()+ ' del aÃ±o ' +str(datetime.now().year))
    adherencia2_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    adherencia2_1.runs[0].font.name = 'Calibri'
    adherencia2_1.runs[0].font.size = Pt(9.5)

    tabla_3 = doc.add_table(rows=9, cols=9)
    tabla_3.style = 'Table Grid'

    def Nombre_mes_actual():
        # Obtener la fecha actual
        fecha_actual = datetime.now()
        # Obtener el nombre del mes
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        if fecha_actual.month == 0:
            mes = meses[11]
        else:
            mes = meses[fecha_actual.month - 1]  # Restamos 1 porque los meses en datetime van de 1 a 12
        return mes

    # Adherencia: contenido en filas
    adherencia_tabla = ["\nMANTENIMIENTO PREVENTIVO\n"]
    adherencia_1 = ["\n\n", "Programado", "Ejecutado", "Programado", "Ejecutado", "Programado", "Ejecutado", "Programado Total", "Ejecutado Total"]
    adherencia_2 = ["CategorÃ­a1 (C1)", "numero", "numero", "-", "-", "-", "-", "numero", "numero"]
    adherencia_3 = ["CategorÃ­a2 (C2)", "-", "-", "numero", "numero", "-", "-", "", ""]
    adherencia_4 = ["CategorÃ­a3 (C3)", "-", "-", "-", "-", "numero", "numero", "", ""]
    adherencia_5 = ["Cumplimiento Semana " + semanas_transcurridas(), "", "", "", "", "", "", "\nporcentaje\n", ""]
    adherencia_6 = ["Cumplimiento de MP planificados en el mes " + Nombre_mes_actual(), "", "", "", "", "", "", "\nfraccion / porcentaje\n", ""]
    adherencia_7 = ["\nActividades de " + Nombre_mes_actual() + " no ejecutadas por Metro\n", "", "", "", "", "", "", "\nfraccion", ""]
    adherencia_8 = ["\nCumplimiento plan Anual (Noviembre 2024-Octubre 2025)\n", "", "", "", "", "", "", "\nfraccion / porcentaje\n", ""]



    # Aplicar formato y contenido
    for i, text in enumerate(adherencia_tabla):
        format_cell(tabla_3.rows[0].cells[i], text, font_name="Calibri", font_size=8, font_bold=False)

    # Agregar datos dinÃ¡micamente con formato
    adherencias = [adherencia_1, adherencia_2, adherencia_3, adherencia_4, adherencia_5, adherencia_6, adherencia_7, adherencia_8]
    for row_idx, adherencia in enumerate(adherencias, start=1):
        for col_idx, text in enumerate(adherencia):
            if text.strip():  # Agregar solo texto no vacÃ­o
                format_cell(tabla_3.rows[row_idx].cells[col_idx], text, font_name="Calibri", font_size=8, font_bold=False)

    tabla_3.rows[2].cells[8].merge(tabla_3.rows[4].cells[8])
    tabla_3.rows[2].cells[7].merge(tabla_3.rows[4].cells[7])

    # FusiÃ³n de celdas adicionales
    for i in range(5, 9):
        tabla_3.rows[i].cells[0].merge(tabla_3.rows[i].cells[5])
        tabla_3.rows[i].cells[6].merge(tabla_3.rows[i].cells[8])
    tabla_3.rows[0].cells[0].merge(tabla_3.rows[0].cells[8])


    tabla_adherencia = doc.add_paragraph('Tabla NÂ°4 â€“ Resumen Mantenimientos Preventivos Semana ' + semanas_transcurridas())
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1
    tabla_adherencia.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    grafico1 = doc.add_paragraph("GrÃ¡fico NÂ°1 â€“ Adherencia Plan de Mantenimiento Preventivo Semana " + semanas_transcurridas() + ".")
    grafico1.alignment = 1
    grafico1.runs[0].font.name = 'Calibri'
    grafico1.runs[0].font.size = Pt(9.5)

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("2.2.2	DesviaciÃ³n al Plan de Mantenimiento", 3)
    plan_mantenimiento.name = 'Calibri'
    plan_mantenimiento.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("       \nA continuaciÃ³n, se detallan las desviaciones al plan de mantenimientos que se han presentado durante el mes en estudio.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    try:
        if not df_filtrado.empty:
            # Filtrar solo las filas donde "TIPO SICE" sea "AverÃ­a"
            df_filtro = df_filtrado[df_filtrado["ESTADO MMS"] == "Emitida"]

            # Definir las columnas requeridas
            columnas_requeridas = ['Descripcion OT', 'DEP', 'EST', 'CAT', 'NÃºmero de OT', 'FE']
            df_filtro2 = df_filtro[columnas_requeridas].copy()

            # Formatear las columnas de fecha
            for col in ['FE', 'FP']:
                if col in df_filtro2.columns:
                    df_filtro2[col] = pd.to_datetime(df_filtro2[col], errors='coerce').dt.date

            # Ordenar las columnas segÃºn el formato deseado
            df_filtro2 = df_filtro2[['Descripcion OT', 'DEP', 'EST', 'CAT', 'NÃºmero de OT', 'FE']]

            # Crear la tabla en el documento de Word
            columnas = ["NÂ°", "DESCRIPCIÃ“N", "LÃ­nea", "EstaciÃ³n", "CAT", "OT", "Plan Matriz", "Motivo DesviaciÃ³n"]
            num_filas = df_filtro2.shape[0]
            num_columnas = len(columnas)
            tabla_desviacion = doc.add_table(rows=num_filas + 1, cols=num_columnas)
            tabla_desviacion.style = 'Table Grid'

            # Agregar encabezados
            for j, column_name in enumerate(columnas):
                cell = tabla_desviacion.cell(0, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Agregar los datos a la tabla
            for i, row in enumerate(df_filtro2.itertuples(index=False), start=1):
                # Agregar numeraciÃ³n en la primera columna
                cell = tabla_desviacion.cell(i, 0)
                cell.text = str(i)
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Agregar el resto de los datos
                for j, value in enumerate(row, start=1):
                    cell = tabla_desviacion.cell(i, j)
                    cell.text = str(value)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Ajustar altura de la fila para la columna de "DESCRIPCIÃ“N"
                    if columnas[j] == "DESCRIPCIÃ“N":
                        cell.width = Inches(3)  # Duplicar el espacio en la celda de descripciÃ³n

                    # Ajustar espacio para la columna "Plan Matriz"
                    if columnas[j] == "Plan Matriz":
                        cell.width = Inches(1.1)

            # Centrar texto en todas las celdas
            for row in tabla_desviacion.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')

    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")


    doc.add_page_break()

    titulo_adherencia2_3 = doc.add_heading('2.2.3 	Adherencia al Plan de Mantenimiento Ãºltimas 12 Semanas \n', 3)
    titulo_adherencia2_3.runs[0].font.name = 'Calibri'
    titulo_adherencia2_3.runs[0].font.size = Pt(10)
    titulo_adherencia2_3.runs[0].font.bold = True
    titulo_adherencia2_3.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    # Crear la tabla
    tabla_4 = doc.add_table(rows=7, cols=9)
    tabla_4.style = 'Table Grid'    

    # Datos de adherencia
    adherencia_tabla = ['\nMANTENIMIENTO PREVENTIVO\n']
    adherencia_1 = ["\n\n", "Programado", "Ejecutado", "Programado", "Ejecutado", "Programado", "Ejecutado", "Programado Total", "Ejecutado Total"]
    adherencia_2 = ["CategorÃ­a1 (C1)", "numero", "numero", "-", "-", "-", "-", "numero", "numero"]
    adherencia_3 = ["CategorÃ­a2 (C2)", "-", "-", "numero", "numero", "-", "-", "", ""]
    adherencia_4 = ["CategorÃ­a3 (C3)", "-", "-", "-", "-", "numero", "numero", "", ""]
    adherencia_5 = ["Cumplimiento ", "", "", "", "", "", "", "\nporcentaje\n", ""]
    adherencia_6 = ["Actividades de Noviembre no ejecutadas por responsabilidad de Metro ", "", "", "", "", "", "", "\nporcentaje\n", ""]

    # FunciÃ³n para dar formato a las celdas con estilo de letra
    def format_cell(cell, text, font_name="Calibri", font_size=8, font_bold=False):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font_name  # Cambiar la fuente
        run.font.size = Pt(font_size)  # Cambiar el tamaÃ±o de letra
        run.font.bold = font_bold  # Negrita
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar horizontalmente
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Aplicar formato y contenido a la primera fila
    for i, text in enumerate(adherencia_tabla):
        format_cell(tabla_4.rows[0].cells[i], text, font_name="Calibri", font_size=8, font_bold=False)

    # Agregar datos dinÃ¡micamente con formato
    adherencias = [adherencia_1, adherencia_2, adherencia_3, adherencia_4, adherencia_5, adherencia_6]
    for row_idx, adherencia in enumerate(adherencias, start=1):
        for col_idx, text in enumerate(adherencia):
            if text.strip():  # Agregar solo texto no vacÃ­o
                format_cell(tabla_4.rows[row_idx].cells[col_idx], text, font_name="Calibri", font_size=8, font_bold=False)

    # FusiÃ³n de celdas
    tabla_4.rows[2].cells[8].merge(tabla_4.rows[4].cells[8])
    tabla_4.rows[2].cells[7].merge(tabla_4.rows[4].cells[7])

    # FusiÃ³n de celdas adicionales
    for i in range(5, 7):
        tabla_4.rows[i].cells[0].merge(tabla_4.rows[i].cells[5])
        tabla_4.rows[i].cells[6].merge(tabla_4.rows[i].cells[8])
    tabla_4.rows[0].cells[0].merge(tabla_4.rows[0].cells[8])

    tabla_adherencia = doc.add_paragraph('Tabla NÂ°6 â€“ Resumen Mantenimientos Preventivos Semana ' + semanas_transcurridas())
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    tabla_adherencia = doc.add_paragraph("GrÃ¡fico NÂ°2 â€“ Adherencia Plan de Mantenimiento Preventivo Semana " + semanas_transcurridas() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("2.2.4	ProyecciÃ³n Mantenimiento PrÃ³ximas 12 Semanas", 3)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se declara proyecciÃ³n del Mantenimiento Preventivo para las prÃ³ximas 12 semanas.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    # ðŸ”¹ Limpiar los datos del DataFrame
    df_tabla12prox.columns = df_tabla12prox.columns.str.strip()  # Limpiar espacios en columnas
    # Limpiar espacios en blanco Ãºnicamente en las columnas que son de tipo texto (str)
    df_tabla12prox.loc[:, df_tabla12prox.select_dtypes(include=["object"]).columns] = (
        df_tabla12prox.select_dtypes(include=["object"]).apply(lambda col: col.str.strip()))
    df_tabla12prox.replace(['nan', 'None', None], '', inplace=True)  # Reemplazar valores nulos con vacÃ­o

    # ðŸ”¹ Convertir la columna `NÂ°` a enteros si existe
    if 'NÂ°' in df_tabla12prox.columns:
        df_tabla12prox['NÂ°'] = pd.to_numeric(df_tabla12prox['NÂ°'], errors='coerce').fillna(0).astype(int)

    # ðŸ”¹ Eliminar filas completamente vacÃ­as
    df_tabla12prox = df_tabla12prox.dropna(how='all').reset_index(drop=True)

    # ðŸ”¹ Verificar el nÃºmero correcto de filas y columnas
    num_filas, num_columnas = df_tabla12prox.shape

    # ðŸ”¹ Crear la tabla en Word
    if num_filas > 0 and num_columnas > 0:  # Asegurarse de que hay datos para insertar
        # Crear la tabla con el nÃºmero exacto de filas
        tabla = doc.add_table(rows=num_filas + 2, cols=num_columnas)  # +2 para tÃ­tulo y encabezado
        tabla.style = 'Table Grid'

        # ðŸ”¹ Fusionar la primera fila y colocar el tÃ­tulo centrado
        titulo_celda = tabla.cell(0, 0)
        titulo_celda.merge(tabla.cell(0, num_columnas - 1))  # Fusionar toda la fila 0
        titulo_celda.text = "ProyecciÃ³n de actividades segÃºn su categorÃ­a"
        titulo_paragraph = titulo_celda.paragraphs[0]
        titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo_paragraph.runs[0]
        titulo_run.font.name = 'Calibri'
        titulo_run.font.size = Pt(10)
        titulo_run.bold = True

        # ðŸ”¹ Insertar encabezados en la fila 1
        for j, column_name in enumerate(df_tabla12prox.columns):
            cell = tabla.cell(1, j)
            cell.text = str(column_name)
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ðŸ”¹ Insertar datos en la tabla desde la fila 2 en adelante
        for i, row in enumerate(df_tabla12prox.itertuples(index=False), start=2):
            for j, value in enumerate(row):
                cell = tabla.cell(i, j)
                cell.text = str(value) if pd.notna(value) else ""
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ðŸ”¹ Fusionar la Ãºltima fila (columna 0 y 1) con el texto "TOTAL"
        ultima_fila = num_filas + 1  # Ãndice de la Ãºltima fila
        cell_total = tabla.cell(ultima_fila, 0)
        cell_total.merge(tabla.cell(ultima_fila, 1))  # Fusionar columnas 0 y 1
        cell_total.text = "TOTAL"
        paragraph = cell_total.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el texto
        run = paragraph.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.bold = True

        tabla_adherencia = doc.add_paragraph('Tabla NÂ°6 â€“ ProyecciÃ³n Mantenimientos Preventivos prÃ³ximas 12 semanas.')
        tabla_adherencia.runs[0].font.name = 'Calibri'
        tabla_adherencia.runs[0].font.size = Pt(9.5)
        tabla_adherencia.runs[0].font.bold = False
        tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("2.3 	Actividades fuera del Plan de Mantenimiento", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se presentan las actividades fuera del plan de Mantenimiento.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    # Crear la tabla con encabezado principal y columnas
    num_columnas = 7  # NÃºmero de columnas segÃºn la imagen
    tabla8 = doc.add_table(rows=2, cols=num_columnas)
    tabla8.style = 'Table Grid'

    # Contenido de las celdas
    encabezado8 = ['\nActividades fuera del plan de Mantenimiento\n']
    sub_encabezado8 = ['NÂ°', 'DescripciÃ³n', 'LÃ­nea', 'EstaciÃ³n', 'Semana', 'Fecha', 'Observaciones']

    # FunciÃ³n para dar formato a las celdas con estilo de letra
    def format_cell(cell, text, font_name="Calibri", font_size=9.5, font_bold= False):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font_name  # Cambiar la fuente
        run.font.size = Pt(font_size)  # Cambiar el tamaÃ±o de letra
        run.font.bold = font_bold  # Negrita
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # AÃ±adir datos y aplicar estilo a cada celda
    for i, text in enumerate(encabezado8):
        format_cell(tabla8.rows[0].cells[i], text, font_name="Calibri", font_size=8, font_bold=True)
    for i, text in enumerate(sub_encabezado8):
        format_cell(tabla8.rows[1].cells[i], text, font_name="Calibri", font_size=8, font_bold =False)

    tabla8.rows[0].cells[0].merge(tabla8.rows[0].cells[6])

    tabla_adherencia = doc.add_paragraph('Tabla NÂ°8 â€“ Actividades fuera del Plan de Mantenimiento Semana ' + semanas_transcurridas() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("3.    Mantenimiento Correctivo.", 1)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    plan_mantenimiento = doc.add_heading("3.1	Fallas Operacionales Semana "+ semanas_transcurridas() + ".", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    plan_mantenimiento = doc.add_heading("3.2	DescripciÃ³n Fallas Operacionales.", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se detallan las averÃ­as presentadas durante la Semana " + semanas_transcurridas() + " del aÃ±o "  + str(datetime.now().year) + ".")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    try:
        if not df_averias.empty:
            # Filtrar solo las filas donde "TIPO SICE" sea "AverÃ­a"
            df_filtro = df_averias[df_averias["TIPO SICE"] == "AverÃ­a"]
            # Definir las columnas requeridas
            columnas_requeridas = ["LINEA", "EMPLAZAMIENTO", "OT", "DESCRIPCIÃ“N DE LA FALLA",
                                   "ACTIVO", "CAT ", "TIPO", "FECHA HORA INFORME", "ESTADO SICE"]
            # Filtrar solo las columnas necesarias
            df_filtro2 =  df_filtro[columnas_requeridas].copy()

            # Formatear la columna 'FECHA HORA INFORME' para que solo contenga la fecha
            if 'FECHA HORA INFORME' in df_filtro2.columns:
                df_filtro2['FECHA HORA INFORME'] = pd.to_datetime(df_filtro2['FECHA HORA INFORME'], errors='coerce').dt.date

            # Crear la tabla en el documento de Word
            num_filas = df_filtro2.shape[0]
            num_columnas = len(columnas_requeridas) + 1  # +1 para la columna "NÂ°"
            tabla = doc.add_table(rows=num_filas + 2, cols=num_columnas)  # +2 para la fila de tÃ­tulo
            tabla.style = 'Table Grid'

            # Agregar la fila inicial con el texto "Fallas Operacionales Semana"
            fila_titulo = tabla.rows[0].cells
            celda_titulo = fila_titulo[0]  # Usar la primera celda para el tÃ­tulo
            celda_titulo.text = "Fallas Operacionales Semana "+semanas_transcurridas()
            # Combinar todas las celdas de la fila
            for celda in fila_titulo[1:]:
                celda_titulo.merge(celda)
            # Formato del tÃ­tulo
            paragraph = celda_titulo.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Calcular el ancho mÃ¡ximo de la columna "DESCRIPCIÃ“N DE LA FALLA"
            max_descripcion_length = df_filtro2["DESCRIPCIÃ“N DE LA FALLA"].astype(str).apply(len).max()
            descripcion_width = Inches(min(6, max(1, max_descripcion_length * 0.1)))  # Ajustar el ancho dinÃ¡micamente

            # Agregar encabezados
            # Columna "NÂ°"
            cell = tabla.cell(1, 0)
            cell.text = "NÂ°"
            paragraph = cell.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Resto de las columnas
            for j, column_name in enumerate(columnas_requeridas, start=1):
                cell = tabla.cell(1, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]

                # Formato del encabezado
                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()

                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Ajustar el ancho de la columna "DESCRIPCIÃ“N DE LA FALLA"
                if column_name == "DESCRIPCIÃ“N DE LA FALLA":
                    cell.width = descripcion_width

            # Agregar los datos a la tabla
            for i, row in enumerate(df_filtro2.itertuples(index=False), start=2):  # start=2 para saltar la fila de tÃ­tulo
                # Columna "NÂ°"
                cell = tabla.cell(i, 0)
                cell.text = str(i - 1)  # Restar 1 para que la numeraciÃ³n comience desde 1
                paragraph = cell.paragraphs[0]
                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Resto de las columnas
                for j, value in enumerate(row, start=1):
                    cell = tabla.cell(i, j)
                    column_name = columnas_requeridas[j - 1]  # Ajustar Ã­ndice de la columna

                    # Formatear la fecha
                    if column_name == "FECHA HORA INFORME" and isinstance(value, (datetime, str)):
                        try:
                            fecha = pd.to_datetime(value).date()
                            cell.text = fecha.strftime("%d-%m-%Y")
                        except:
                            cell.text = str(value)
                    else:
                        cell.text = str(value) if pd.notna(value) else ""

                    paragraph = cell.paragraphs[0]

                    if paragraph.runs:
                        run = paragraph.runs[0]
                    else:
                        run = paragraph.add_run()

                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Ajustar espacio para columnas especÃ­ficas
                    if column_name == "DESCRIPCIÃ“N DE LA FALLA":
                        cell.width = descripcion_width
                    elif column_name == "FECHA HORA INFORME":
                        cell.width = Inches(1.5)

            # Centrar texto en todas las celdas
            for row in tabla.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')
    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")


#############################
    tabla_adherencia = doc.add_paragraph('Tabla NÂ°9 â€“ Fallas Operacionales Semana ' + semanas_transcurridas() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("3.3	Comportamiento fallas operacionales Ãºltimas 12 semanas. ", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se grÃ¡fica el comportamiento de las fallas operacionales de las Ãºltimas 12 semanas segÃºn sus categorÃ­as.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    tabla_adherencia = doc.add_paragraph("GrÃ¡fico NÂ°3 â€“ Comportamiento fallas operacionales Ãºltimas 12 semanas categorÃ­a 1.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    tabla_adherencia = doc.add_paragraph("GrÃ¡fico NÂ°4 â€“ Comportamiento fallas operacionales Ãºltimas 12 semanas categorÃ­a 2.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    tabla_adherencia = doc.add_paragraph("GrÃ¡fico NÂ°5 â€“ Comportamiento fallas operacionales Ãºltimas 12 semanas categorÃ­a 3.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("3.4 	ComparaciÃ³n fallas operacionales. ", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se presenta la comparaciÃ³n 2021-2022-2023-2024-2025 de las fallas operacionales del sistema de comunicaciones segÃºn su categorÃ­a. ")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    tabla_adherencia = doc.add_paragraph("\n    GrÃ¡fico NÂ°6 â€“ Comportamiento de las fallas operacionales en los Ãºltimos 4 aÃ±os categorÃ­as 1, 2 y 3.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    tabla_adherencia = doc.add_paragraph("\n    GrÃ¡fico NÂ°7 â€“ Comportamiento mes a mes de las fallas operacionales en los Ãºltimos 4 aÃ±os categorÃ­a 1.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    tabla_adherencia = doc.add_paragraph("\n    GrÃ¡fico NÂ°8 â€“ Comportamiento mes a mes de las fallas operacionales en los Ãºltimos 4 aÃ±os categorÃ­a 2.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    tabla_adherencia = doc.add_paragraph("\n    GrÃ¡fico NÂ°9 â€“ Comportamiento mes a mes de las fallas operacionales en los Ãºltimos 4 aÃ±os categorÃ­a 3.")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("3.5 	Resumen fallas operacionales del mes en curso.", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      Para la semana en estudio, durante el mes de diciembre se ha informado un total 15 averÃ­as asociadas al sistema de comunicaciones.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)
    
    try:
        if not df_averias.empty:
            # Filtrar solo las filas donde "TIPO SICE" sea "AverÃ­a"
            df_filtro = df_averias_mes[df_averias_mes["TIPO SICE"] == "AverÃ­a"]
            # Definir las columnas requeridas
            columnas_requeridas = ["LINEA", "EMPLAZAMIENTO", "OT", "DESCRIPCIÃ“N DE LA FALLA",
                                   "ACTIVO", "CAT ", "TIPO", "FECHA HORA INFORME", "ESTADO SICE", "Semana"]
            # Filtrar solo las columnas necesarias
            df_filtro2 = df_filtro[columnas_requeridas].copy()

            # Formatear la columna 'FECHA HORA INFORME' para que solo contenga la fecha
            if 'FECHA HORA INFORME' in df_filtro2.columns:
                df_filtro2['FECHA HORA INFORME'] = pd.to_datetime(df_filtro2['FECHA HORA INFORME'], errors='coerce').dt.date

            # Crear la tabla en el documento de Word
            num_filas = df_filtro2.shape[0]
            num_columnas = len(columnas_requeridas) + 1  # +1 para la columna "NÂ°"
            tabla10 = doc.add_table(rows=num_filas + 2, cols=num_columnas)  # +2 para la fila de tÃ­tulo
            tabla10.style = 'Table Grid'

            # Agregar la fila inicial con el texto "Fallas Operacionales Semana"
            fila_titulo = tabla10.rows[0].cells
            celda_titulo = fila_titulo[0]  # Usar la primera celda para el tÃ­tulo
            celda_titulo.text = "Fallas Operacionales mes de " + Nombre_mes_actual()
            # Combinar todas las celdas de la fila
            for celda in fila_titulo[1:]:
                celda_titulo.merge(celda)
            # Formato del tÃ­tulo
            paragraph = celda_titulo.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Calcular el ancho mÃ¡ximo de la columna "DESCRIPCIÃ“N DE LA FALLA"
            max_descripcion_length = df_filtro2["DESCRIPCIÃ“N DE LA FALLA"].astype(str).apply(len).max()
            descripcion_width = Inches(min(6, max(1, max_descripcion_length * 0.1)))  # Ajustar el ancho dinÃ¡micamente

            # Agregar encabezados
            # Columna "NÂ°"
            cell = tabla10.cell(1, 0)
            cell.text = "NÂ°"
            paragraph = cell.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Resto de las columnas
            for j, column_name in enumerate(columnas_requeridas, start=1):
                cell = tabla10.cell(1, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]

                # Formato del encabezado
                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()

                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Ajustar el ancho de la columna "DESCRIPCIÃ“N DE LA FALLA"
                if column_name == "DESCRIPCIÃ“N DE LA FALLA":
                    cell.width = descripcion_width

            # Agregar los datos a la tabla10
            for i, row in enumerate(df_filtro2.itertuples(index=False), start=2):  # start=2 para saltar la fila de tÃ­tulo
                # Columna "NÂ°"
                cell = tabla10.cell(i, 0)
                cell.text = str(i - 1)  # Restar 1 para que la numeraciÃ³n comience desde 1
                paragraph = cell.paragraphs[0]
                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Resto de las columnas
                for j, value in enumerate(row, start=1):
                    cell = tabla10.cell(i, j)
                    column_name = columnas_requeridas[j - 1]  # Ajustar Ã­ndice de la columna

                    # Formatear la fecha
                    if column_name == "FECHA HORA INFORME" and isinstance(value, (datetime, str)):
                        try:
                            fecha = pd.to_datetime(value).date()
                            cell.text = fecha.strftime("%d-%m-%Y")
                        except:
                            cell.text = str(value)
                    else:
                        cell.text = str(value) if pd.notna(value) else ""

                    paragraph = cell.paragraphs[0]

                    if paragraph.runs:
                        run = paragraph.runs[0]
                    else:
                        run = paragraph.add_run()

                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Ajustar espacio para columnas especÃ­ficas
                    if column_name == "DESCRIPCIÃ“N DE LA FALLA":
                        cell.width = descripcion_width
                    elif column_name == "FECHA HORA INFORME":
                        cell.width = Inches(1.5)

            # Centrar texto en todas las celdas
            for row in tabla10.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')
    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")

    tabla_adherencia = doc.add_paragraph("\n    Tabla NÂ°10â€“ AnÃ¡lisis de fallas operacionales mes de " + Nombre_mes_actual() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("3.6	Resumen fallas operacionales escaladas a proveedores", 2)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      Durante la Semana " + semanas_transcurridas() + " no se informan averÃ­as escaladas a proveedores para anÃ¡lisis y soluciÃ³n.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    # Crear la tabla con encabezado principal y columnas
    num_columnas = 9  # NÃºmero de columnas segÃºn la imagen
    tabla11 = doc.add_table(rows=3, cols=num_columnas)
    tabla11.style = 'Table Grid'

    # Contenido de las celdas
    encabezado8 = ["\nFallas Operacionales Escaladas a Proveedores en la Semana " + semanas_transcurridas() +".\n"]
    sub_encabezado8 = ['NÂ°', 'LÃ­nea', 'EstaciÃ³n', 'OT', 'DescripciÃ³n de la Falla', 'CategorÃ­a', 'Estado', 'Semana', 'Proveedor']

    # FunciÃ³n para dar formato a las celdas con estilo de letra
    def format_cell(cell, text, font_name="Calibri", font_size=9.5, font_bold= False):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font_name  # Cambiar la fuente
        run.font.size = Pt(font_size)  # Cambiar el tamaÃ±o de letra
        run.font.bold = font_bold  # Negrita
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # AÃ±adir datos y aplicar estilo a cada celda
    for i, text in enumerate(encabezado8):
        format_cell(tabla11.rows[0].cells[i], text, font_name="Calibri", font_size=6, font_bold=True)
    for i, text in enumerate(sub_encabezado8):
        format_cell(tabla11.rows[1].cells[i], text, font_name="Calibri", font_size=6, font_bold =False)

    tabla11.rows[0].cells[0].merge(tabla11.rows[0].cells[8])

    tabla_adherencia = doc.add_paragraph('Tabla NÂ°11â€“ AnÃ¡lisis Falla Operacionales Semana ' + semanas_transcurridas() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("4     Componente extraordinaria ", 1)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se detallan las componentes extraordinarias presentadas en el mes.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    try:
        if not df_averias.empty:
            # Filtrar solo las filas donde "TIPO SICE" sea "Vandalismo"
            df_filtro = df_averias_mes[df_averias_mes["TIPO SICE"] == "Vandalismo"]
            # Definir las columnas requeridas
            columnas_requeridas = ["LINEA", "EMPLAZAMIENTO", "OT", "DESCRIPCIÃ“N DE LA FALLA",
                                   "SISTEMA", "CAT ", "TIPO", "FECHA HORA INFORME", "Semana"]
            # Filtrar solo las columnas necesarias
            df_filtro2 = df_filtro[columnas_requeridas].copy()

            # Modificar los nombres de las columnas
            df_filtro2.columns = ["LINEA", "EMPLAZAMIENTO", "OT", "DESCRIPCIÃ“N DE LA FALLA",
                                  "SISTEMA", "Autorizado", "CÃ³digo valor", "Fecha de soluciÃ³n", "Semana"]

            # Formatear la columna 'FECHA HORA INFORME' para que solo contenga la fecha
            if 'Fecha de soluciÃ³n' in df_filtro2.columns:
                df_filtro2['Fecha de soluciÃ³n'] = pd.to_datetime(df_filtro2['Fecha de soluciÃ³n'], errors='coerce').dt.date

            # Crear la tabla en el documento de Word
            num_filas = df_filtro2.shape[0]
            num_columnas = len(columnas_requeridas) + 1  # +1 para la columna "NÂ°"
            tabla10 = doc.add_table(rows=num_filas + 2, cols=num_columnas)  # +2 para la fila de tÃ­tulo
            tabla10.style = 'Table Grid'

            # Agregar la fila inicial con el texto "Fallas Operacionales Semana"
            fila_titulo = tabla10.rows[0].cells
            celda_titulo = fila_titulo[0]  # Usar la primera celda para el tÃ­tulo
            celda_titulo.text = "Componente Extraordinaria."
            # Combinar todas las celdas de la fila
            for celda in fila_titulo[1:]:
                celda_titulo.merge(celda)
            # Formato del tÃ­tulo
            paragraph = celda_titulo.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Calcular el ancho mÃ¡ximo de la columna "DESCRIPCIÃ“N DE LA FALLA"
            max_descripcion_length = df_filtro2["DESCRIPCIÃ“N DE LA FALLA"].astype(str).apply(len).max()
            descripcion_width = Inches(min(6, max(1, max_descripcion_length * 0.1)))  # Ajustar el ancho dinÃ¡micamente

            # Agregar encabezados
            # Columna "NÂ°"
            cell = tabla10.cell(1, 0)
            cell.text = "NÂ°"
            paragraph = cell.paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
            else:
                run = paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Resto de las columnas
            for j, column_name in enumerate(df_filtro2.columns, start=1):
                cell = tabla10.cell(1, j)
                cell.text = column_name
                paragraph = cell.paragraphs[0]

                # Formato del encabezado
                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()

                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                run.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Ajustar el ancho de la columna "DESCRIPCIÃ“N DE LA FALLA"
                if column_name == "DESCRIPCIÃ“N DE LA FALLA":
                    cell.width = descripcion_width

            # Agregar los datos a la tabla10
            for i, row in enumerate(df_filtro2.itertuples(index=False), start=2):  # start=2 para saltar la fila de tÃ­tulo
                # Columna "NÂ°"
                cell = tabla10.cell(i, 0)
                cell.text = str(i - 1)  # Restar 1 para que la numeraciÃ³n comience desde 1
                paragraph = cell.paragraphs[0]
                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Resto de las columnas
                for j, value in enumerate(row, start=1):
                    cell = tabla10.cell(i, j)
                    column_name = df_filtro2.columns[j - 1]  # Ajustar Ã­ndice de la columna

                    # Formatear la fecha
                    if column_name == "Fecha de soluciÃ³n" and isinstance(value, (datetime, str)):
                        try:
                            fecha = pd.to_datetime(value).date()
                            cell.text = fecha.strftime("%d-%m-%Y")
                        except:
                            cell.text = str(value)
                    else:
                        cell.text = str(value) if pd.notna(value) else ""

                    paragraph = cell.paragraphs[0]

                    if paragraph.runs:
                        run = paragraph.runs[0]
                    else:
                        run = paragraph.add_run()

                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Ajustar espacio para columnas especÃ­ficas
                    if column_name == "DESCRIPCIÃ“N DE LA FALLA":
                        cell.width = descripcion_width
                    elif column_name == "Fecha de soluciÃ³n":
                        cell.width = Inches(1.5)

            # Centrar texto en todas las celdas
            for row in tabla10.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            doc.add_paragraph('\nNo se encontraron datos para mostrar.')
    except Exception as e:
        print(f"Error en la creaciÃ³n del Word: {e}")

    doc.add_page_break()

    plan_mantenimiento = doc.add_heading("5     GestiÃ³n INFOR ", 1)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, se muestra reporte del estado de las Ã³rdenes de trabajo.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    # Crear la tabla con encabezado principal y columnas
    num_columnas = 7  # NÃºmero de columnas segÃºn la imagen
    tabla13 = doc.add_table(rows=5, cols=num_columnas)
    tabla13.style = 'Table Grid'

    # Contenido de las celdas
    encabezado8 = ["\nEstado de OT en Sistema INFOR\n"]
    sub_encabezado8 = ['Estado', 'En Curso', 'Emitida', 'Pendiente', 'En VerificaciÃ³n', 'Completada', 'Cerrada']

    # FunciÃ³n para dar formato a las celdas con estilo de letra
    def format_cell(cell, text, font_name="Calibri", font_size=9.5, font_bold= False):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font_name  # Cambiar la fuente
        run.font.size = Pt(font_size)  # Cambiar el tamaÃ±o de letra
        run.font.bold = font_bold  # Negrita
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # AÃ±adir datos y aplicar estilo a cada celda
    for i, text in enumerate(encabezado8):
        format_cell(tabla13.rows[0].cells[i], text, font_name="Calibri", font_size=8, font_bold=True)
    for i, text in enumerate(sub_encabezado8):
        format_cell(tabla13.rows[1].cells[i], text, font_name="Calibri", font_size=8, font_bold =False)

    tabla13.rows[0].cells[0].merge(tabla13.rows[0].cells[6])

    # Llenar columnas 3, 4 y 5 con los datos proporcionados
    datos = [
        ["Vandalismos", "0", "0", "0", "0", "0", "0"],
        ["Semana " + semanas_transcurridas(), "0", "0", "0", "0", "0", "0"],
        ["HistÃ³rico " + str(datetime.now().year), "0", "0", "0", "0", "0", "0"]]

    for row_idx, row_data in enumerate(datos, start=2):  # Empezar desde la fila 2
        for col_idx, cell_data in enumerate(row_data):
            format_cell(tabla13.rows[row_idx].cells[col_idx], cell_data, font_name="Calibri", font_size=8, font_bold=False)

    tabla_adherencia = doc.add_paragraph('Tabla NÂ°13 â€“ Tabla de gestiÃ³n de INFOR Semana ' + Nombre_mes_actual() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1

    doc.add_page_break()
    
    plan_mantenimiento = doc.add_heading("6     Hallazgos de seguridad ", 1)
    plan_mantenimiento.runs[0].font.name = 'Calibri'
    plan_mantenimiento.runs[0].font.size = Pt(10)
    plan_mantenimiento.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Negro

    mantenimiento3 = doc.add_paragraph("\n      A continuaciÃ³n, los hallazgos de seguridad de la semana en estudio.")
    mantenimiento3.runs[0].font.name = 'Calibri'
    mantenimiento3.runs[0].font.size = Pt(9.5)

    # Crear la tabla con encabezado principal y columnas
    num_columnas = 3  # NÃºmero de columnas segÃºn la imagen
    tabla13 = doc.add_table(rows=2, cols=num_columnas)
    tabla13.style = 'Table Grid'

    # Contenido de las celdas
    sub_encabezado8 = ['Hallazgo', 'Plan de acciÃ³n', 'Fecha de soluciÃ³n']

    # FunciÃ³n para dar formato a las celdas con estilo de letra
    def format_cell(cell, text, font_name="Calibri", font_size=9.5, font_bold= False):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font_name  # Cambiar la fuente
        run.font.size = Pt(font_size)  # Cambiar el tamaÃ±o de letra
        run.font.bold = font_bold  # Negrita
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # AÃ±adir datos y aplicar estilo a cada celda
    for i, text in enumerate(sub_encabezado8):
        format_cell(tabla13.rows[0].cells[i], text, font_name="Calibri", font_size=8, font_bold =False)

    tabla_adherencia = doc.add_paragraph('Tabla NÂ°14 â€“ Hallazgos de seguridad Semana ' + semanas_transcurridas() + ".")
    tabla_adherencia.runs[0].font.name = 'Calibri'
    tabla_adherencia.runs[0].font.size = Pt(9.5)
    tabla_adherencia.runs[0].font.bold = False
    tabla_adherencia.alignment = 1
    


################################# 
    # Abrir un cuadro de diÃ¡logo para seleccionar la ruta de guardado
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")])

    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Guardado exitoso", f"Archivo guardado correctamente en {file_path}")
    else:
        messagebox.showwarning("Guardado cancelado", "El guardado del archivo ha sido cancelado")
    print(f"Portada agregada correctamente. Archivo guardado")

# Variable global para la lÃ³gica (Y u O)
logica_filtro = tk.StringVar(value="Y")  # Valor inicial: "Y"

# Configurar el frame_independiente para tener espacio adicional para el botÃ³n "Y/O"
total_columnas_descripcion = 5  # Aumentar a 5 para incluir espacio adicional para el botÃ³n "Y/O"

# Configurar las columnas para centrar el contenido
for col in range(total_columnas_descripcion):
    if col == 0 or col == total_columnas_descripcion - 1:  # Columnas vacÃ­as (izquierda y derecha)
        frame_independiente.grid_columnconfigure(col, weight=1)  # Espacio expansible
    else:  # Columnas con contenido
        frame_independiente.grid_columnconfigure(col, weight=0)  # Fijas

# Mover el botÃ³n "Y/O" a la derecha del frame_independiente
label_logica = tk.Label(
    frame_independiente,
    text="LÃ³gica (Y/O)",
    bg="#2b2b2b",
    fg="white",
    font=("arial", 10)
)
label_logica.grid(row=0, column=3, padx=10, pady=5, sticky="e")  # Columna 3 para el texto "LÃ³gica (Y/O)"

menu_logica = tk.OptionMenu(
    frame_independiente,
    logica_filtro,
    "Y", "O",
    command=actualizar_logica
)
menu_logica.config(width=10, bg="white", fg="#2b2b2b")
menu_logica.grid(row=0, column=4, padx=10, pady=5, sticky="w")  # Columna 4 para el menÃº desplegable

# BotÃ³n para exportar datos seleccionados
boton_exportar = tk.Button(
    frame_anexo,
    text="Exportar Datos Seleccionados",
    command=exportar_seleccion,
    font=("arial", 10, "bold"),
    activebackground="#4c70ba",
    activeforeground="white",
    bd=2,
    relief="raised"
)
boton_exportar.grid(row=5, column=0, pady=10, sticky="n")

# BotÃ³n para volver al menÃº principal
boton_volver = tk.Button(
    frame_anexo, text="â† Volver al MenÃº",
    command=lambda: cambiar_frame(frame_anexo, frame_menu),
    font=("arial", 10)
)
boton_volver.grid(pady=10)

# Crear el frame_semanal (nuevo frame)
frame_semanal = tk.Frame(ventana, bg="#2b2b2b")
frame_menu.pack(fill="both", expand=True)

label_semanal = tk.Label(frame_semanal, text="Informe Semanal", bg="#2b2b2b", fg="white", font=("arial", 20))
label_semanal.pack(pady=20)

# Crear un frame para organizar los botones horizontalmente
frame_botones = tk.Frame(frame_semanal, bg="#2b2b2b")
frame_botones.pack(pady=10)

# BotÃ³n para abrir archivo Excel
boton_abrir_programacion = tk.Button(
    frame_botones, text="Abrir Archivo Programacion",
    command=abrir_programacion,
    font=("arial", 10)
)
boton_abrir_programacion.pack(side="left", padx=10, pady=5)

# ðŸ”¹ Nuevo botÃ³n al lateral derecho de abrir_programacion
boton_abrir_averias = tk.Button(
    frame_botones, text="Abrir Archivo Averias",
    command=abrir_averias,
    font=("arial", 10)
)
boton_abrir_averias.pack(side="left", padx=10, pady=5)  # ðŸ”¹ Alineado a la derecha del anterior


# AquÃ­ se agrega el menÃº desplegable al frame_semanal
opciones = ["A", "B", "0"]
Rev_documento = tk.StringVar(value="REV. documento")  # Variable para almacenar la selecciÃ³n

menu_opciones = tk.OptionMenu(
    frame_semanal,
    Rev_documento,
    *opciones
)
menu_opciones.config(
    font=("arial", 8),
    activebackground="#4b4b4b",
    activeforeground="white"
)
menu_opciones["menu"].config(bg="#1e1e1e", fg="white")  # Estilo del menÃº desplegable
menu_opciones.pack(pady=10)

# BotÃ³n para exportar datos seleccionados
boton_exportar = tk.Button(
    frame_semanal,
    text="Exportar Datos Seleccionados",
    command=crear_word,
    font=("arial", 10, "bold"),
    activebackground="#4c70ba",
    activeforeground="white",
    bd=2,
    relief="raised"
)
boton_exportar.pack(pady=10)

# BotÃ³n para volver al menÃº principal
boton_volver_nuevo = tk.Button(
    frame_semanal, text="â† Volver al MenÃº",
    command=lambda: cambiar_frame(frame_semanal, frame_menu),
    font=("arial", 10)
)
boton_volver_nuevo.pack(pady=10)

# Botones del frame_menu
boton_anexo.bind("<Enter>", on_enter)
boton_anexo.bind("<Leave>", on_leave)

boton_nuevo.bind("<Enter>", on_enter)
boton_nuevo.bind("<Leave>", on_leave)

# Botones del frame_anexo
boton_insertar.bind("<Enter>", on_enter)
boton_insertar.bind("<Leave>", on_leave)

boton_exportar.bind("<Enter>", on_enter)
boton_exportar.bind("<Leave>", on_leave)

boton_volver_nuevo.bind("<Enter>", on_enter)
boton_volver_nuevo.bind("<Leave>", on_leave)

boton_volver.bind("<Enter>", on_enter)
boton_volver.bind("<Leave>", on_leave)

# Botones del frame_semanal
boton_abrir_programacion.bind("<Enter>", on_enter)
boton_abrir_programacion.bind("<Leave>", on_leave)

boton_volver_nuevo.bind("<Enter>", on_enter)
boton_volver_nuevo.bind("<Leave>", on_leave)

boton_abrir_averias.bind("<Enter>", on_enter)
boton_abrir_averias.bind("<Leave>", on_leave)

# Iniciar la aplicaciÃ³n
ventana.mainloop()